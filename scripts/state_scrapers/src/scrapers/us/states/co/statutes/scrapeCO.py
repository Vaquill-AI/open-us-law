"""Colorado Revised Statutes (CRS) scraper.

Source: Wayback Machine mirror of official CRS DOCX files published by the
Colorado Office of Legislative Legal Services (OLLS). The canonical origin is
https://leg.colorado.gov/sites/default/files/images/olls/crs2024-title-NN.docx
but that URL returns HTTP 403 from all IPs. The identical DOCX is archived at
https://web.archive.org/web/2024/https://leg.colorado.gov/sites/default/files/images/olls/crs2024-title-NN.docx
and returns 200 without any auth requirement.

Hierarchy: corpus -> title -> article -> [part ->] section
Section text is embedded inline in each DOCX.

Section citation format: "C.R.S. § TITLE-ARTICLE-SECTION"
e.g. C.R.S. § 1-1-101
"""
from __future__ import annotations

import os
import re
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Project-root bootstrap (mirrors ME/ND pattern)
# ---------------------------------------------------------------------------
_current_file = Path(__file__).resolve()
_src_dir = _current_file.parent
while _src_dir.name != "src" and _src_dir.parent != _src_dir:
    _src_dir = _src_dir.parent
_project_root = _src_dir.parent
if str(_project_root) not in sys.path:
    sys.path.append(str(_project_root))

from src.utils.pydanticModels import Addendum, AddendumType, Node, NodeText
from src.utils.scrapingHelpers import (
    insert_jurisdiction_and_corpus_node,
    insert_node,
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
COUNTRY = "us"
JURISDICTION = "co"
CORPUS = "statutes"
TABLE_NAME = f"{COUNTRY}_{JURISDICTION}_{CORPUS}"

# Wayback Machine base URL for the DOCX archive. The `2024` date-range glob
# resolves to the most recent snapshot from that year automatically.
_WAYBACK_BASE = (
    "https://web.archive.org/web/2024/"
    "https://leg.colorado.gov/sites/default/files/images/olls/"
)
# Canonical origin URL (403 directly, used only as the section link target).
_ORIGIN_BASE = (
    "https://leg.colorado.gov/sites/default/files/images/olls/"
)
# Year prefix in file name.
_CRS_YEAR = "2024"

# All CRS titles: integers 1-44 plus the three decimal titles.
# The order matches the official CRS publication order.
_ALL_TITLES = [
    "01", "02", "03", "04", "05", "06", "07", "08", "09", "10",
    "11", "12", "13", "14", "15", "16", "17", "18", "19", "20",
    "21", "22", "23", "24", "25", "25.5", "26", "26.5", "27", "28",
    "29", "30", "31", "32", "33", "34", "35", "36", "37", "38",
    "39", "40", "41", "42", "43", "44",
]

_RESERVED_KEYWORDS = ["(repealed)", "(expired)", "(reserved)", "(renumbered)"]

# Regex: section header line.
# Examples:  "1-1-101.  Short title..."
#            "25.5-1-102.  Legislative declaration..."
#            "1-7.5-108.  Mail-in ballots. (Repealed)"
#            "1-1-401 to 1-1-403. (Repealed)"  (range notation)
_SEC_RE = re.compile(
    r"^(\d{1,2}(?:\.\d+)?-\d+(?:\.\d+)?-\d+(?:\.\d+)?)(?:\s+to\s+[\d\.\-]+)?\.\s*(.*)"
)

# Browser-like headers for Wayback Machine (avoids bot detection on .archive.org).
_REQ_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/131.0.0.0 Safari/537.36"
    ),
    "Accept": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document,"
        "application/octet-stream,*/*"
    ),
    "Referer": "https://web.archive.org/",
}

# Maximum number of Wayback fetch retries per title.
_MAX_RETRIES = 3


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    corpus_node: Node = insert_jurisdiction_and_corpus_node(COUNTRY, JURISDICTION, CORPUS)
    _scrape_all_titles(corpus_node)


def _scrape_all_titles(corpus_node: Node) -> None:
    """Dispatch every CRS title in parallel.

    The title list is static (_ALL_TITLES), so there is no discovery pass:
    each title is one self-contained DOCX download plus an in-memory parse
    that shares no state with the other titles, and the JSONL sink + counters
    in vaquill_pipeline.patch are lock-protected. Titles therefore fan out to
    a ThreadPoolExecutor sized by VAQUILL_TITLE_WORKERS (default 8, matching
    scrapeWA / scrapeAK).

    Why: CO is a small number of large requests (~45 DOCX files off the
    Wayback Machine), so the crawl is dominated by download + parse latency
    per title rather than by request count. Overlapping the titles collapses
    that wall-clock without raising the request count at all.

    NOTE: CO deliberately has no titles_done resume file -- every run
    re-downloads and re-parses each title in full. That is what lets an
    amended section be re-chunked into a fresh content-addressed point_id
    (the JSONL skipset suppresses the write for unchanged sections, so a
    re-run is cheap in output but still catches amendments). Do not add a
    titles_done skip here to save time without replacing that freshness some
    other way -- it would make CO amendment-blind.
    """

    def _do_title(title_slug: str) -> tuple[str, str, str | None]:
        # One title's failure must not abort the other workers, so each is
        # wrapped and reported; the run continues with the remaining titles.
        try:
            _scrape_title(corpus_node, title_slug)
            return (title_slug, "ok", None)
        except Exception as exc:
            return (title_slug, "fail", str(exc)[:200])

    workers = int(os.environ.get("VAQUILL_TITLE_WORKERS", "8"))
    print(
        f"[scrapeCO] running {len(_ALL_TITLES)} titles with {workers} parallel workers",
        flush=True,
    )
    with ThreadPoolExecutor(max_workers=workers) as ex:
        for fut in as_completed(ex.submit(_do_title, item) for item in _ALL_TITLES):
            slug, status, err = fut.result()
            if status == "fail":
                print(f"[scrapeCO] title {slug}: {status}: {err}", flush=True)
            else:
                print(f"[scrapeCO] title {slug}: {status}", flush=True)


# ---------------------------------------------------------------------------
# Title level
# ---------------------------------------------------------------------------

def _scrape_title(corpus_node: Node, title_slug: str) -> None:
    """Download one title DOCX and parse all structure + content nodes."""
    docx_bytes = _fetch_docx(title_slug)
    if docx_bytes is None:
        print(f"[CO] SKIP title={title_slug}: could not download DOCX", flush=True)
        return

    # The canonical top_level_title is the numeric label (e.g. "1", "25.5").
    top_level_title = title_slug.lstrip("0") or "0"

    origin_url = f"{_ORIGIN_BASE}crs{_CRS_YEAR}-title-{title_slug}.docx"

    paragraphs = _extract_paragraphs(docx_bytes)
    if not paragraphs:
        print(f"[CO] SKIP title={title_slug}: no paragraphs extracted", flush=True)
        return

    # Locate the TITLE header and derive the title name.
    title_name, title_start = _find_title_header(paragraphs, top_level_title)
    if title_start < 0:
        print(f"[CO] WARN title={title_slug}: could not find TITLE header", flush=True)
        title_start = 0
        title_name = f"Title {top_level_title}"

    node_id = f"{corpus_node.node_id}/title={top_level_title}"
    title_node = Node(
        id=node_id,
        link=origin_url,
        top_level_title=top_level_title,
        node_type="structure",
        level_classifier="title",
        number=top_level_title,
        node_name=title_name,
        parent=corpus_node.node_id,
    )
    insert_node(title_node, TABLE_NAME, ignore_duplicate=True, debug_mode=True)

    _parse_title_body(title_node, paragraphs[title_start:], origin_url)


# ---------------------------------------------------------------------------
# Body parser: walks paragraphs and emits article / part / section nodes
# ---------------------------------------------------------------------------

def _parse_title_body(
    title_node: Node,
    paragraphs: list[str],
    origin_url: str,
) -> None:
    """
    State-machine parser over a title's paragraph list.

    State stack:
        article_node  -- current ARTICLE (required)
        part_node     -- current PART (optional; some articles have no parts)
    """
    current_article: Optional[Node] = None
    current_part: Optional[Node] = None
    # Accumulate pending lines for the next article/part name.
    pending_label: list[str] = []
    pending_type: Optional[str] = None  # "ARTICLE" or "PART"

    # Current section accumulation
    cur_sec_number: Optional[str] = None
    cur_sec_name: Optional[str] = None
    cur_sec_lines: list[str] = []
    cur_sec_parent: Optional[Node] = None

    def flush_section() -> None:
        nonlocal cur_sec_number, cur_sec_name, cur_sec_lines, cur_sec_parent
        if cur_sec_number is None or cur_sec_parent is None:
            return
        _emit_section(
            cur_sec_number,
            cur_sec_name or cur_sec_number,
            cur_sec_lines,
            cur_sec_parent,
            origin_url,
        )
        cur_sec_number = None
        cur_sec_name = None
        cur_sec_lines = []
        cur_sec_parent = None

    i = 0
    while i < len(paragraphs):
        raw = paragraphs[i]
        text = raw.strip()
        i += 1

        if not text:
            continue

        # ------------------------------------------------------------------
        # ARTICLE header line: "ARTICLE N"
        # ------------------------------------------------------------------
        m_art = re.match(r"^ARTICLE\s+(\d+(?:\.\d+)?)$", text)
        if m_art:
            flush_section()
            art_number = m_art.group(1)
            # Collect the name from the NEXT non-empty line(s).
            art_name_parts: list[str] = [f"Article {art_number}"]
            j = i
            while j < len(paragraphs):
                nxt = paragraphs[j].strip()
                j += 1
                if not nxt:
                    continue
                # Stop if the next line is another structural marker.
                if _is_structural(nxt):
                    i = j - 1
                    break
                art_name_parts.append(nxt)
                # Check if consecutive lines together form the full name.
                # A single descriptive line suffices.
                # Two-line article names exist (e.g., "GENERAL, PRIMARY,\nCONGRESSIONAL VACANCY").
                # Stop after collecting the first descriptive line that does not
                # continue with uppercase-only text.
                if not nxt.isupper():
                    i = j
                    break
            else:
                i = j

            full_art_name = " ".join(art_name_parts)
            art_node_id = f"{title_node.node_id}/article={art_number}"
            current_article = Node(
                id=art_node_id,
                link=origin_url,
                top_level_title=title_node.top_level_title,
                node_type="structure",
                level_classifier="article",
                number=art_number,
                node_name=full_art_name,
                parent=title_node.node_id,
            )
            insert_node(current_article, TABLE_NAME, ignore_duplicate=True, debug_mode=True)
            current_part = None
            continue

        # ------------------------------------------------------------------
        # PART header line: "PART N"
        # ------------------------------------------------------------------
        m_part = re.match(r"^PART\s+(\d+(?:\.\d+)?)$", text)
        if m_part:
            flush_section()
            part_number = m_part.group(1)
            art_parent = current_article or title_node

            part_name_parts: list[str] = [f"Part {part_number}"]
            j = i
            while j < len(paragraphs):
                nxt = paragraphs[j].strip()
                j += 1
                if not nxt:
                    continue
                if _is_structural(nxt):
                    i = j - 1
                    break
                part_name_parts.append(nxt)
                if not nxt.isupper():
                    i = j
                    break
            else:
                i = j

            full_part_name = " ".join(part_name_parts)
            part_node_id = f"{art_parent.node_id}/part={part_number}"
            current_part = Node(
                id=part_node_id,
                link=origin_url,
                top_level_title=title_node.top_level_title,
                node_type="structure",
                level_classifier="part",
                number=part_number,
                node_name=full_part_name,
                parent=art_parent.node_id,
            )
            insert_node(current_part, TABLE_NAME, ignore_duplicate=True, debug_mode=True)
            continue

        # ------------------------------------------------------------------
        # Section header line
        # ------------------------------------------------------------------
        m_sec = _SEC_RE.match(text)
        if m_sec:
            flush_section()
            sec_number = m_sec.group(1)
            sec_rest = m_sec.group(2).strip()
            # Derive human-readable section name from the rest of the line.
            # The rest starts with: "Short title. Body text..." or "(Repealed)"
            sec_name = _extract_section_name(sec_rest)
            sec_parent = current_part or current_article or title_node
            cur_sec_number = sec_number
            cur_sec_name = sec_name
            cur_sec_lines = [text]
            cur_sec_parent = sec_parent
            continue

        # ------------------------------------------------------------------
        # Continuation of current section
        # ------------------------------------------------------------------
        if cur_sec_number is not None:
            cur_sec_lines.append(text)
            continue

        # Other lines before any section starts: skip (article name, notes, etc.)

    flush_section()


# ---------------------------------------------------------------------------
# Section emitter
# ---------------------------------------------------------------------------

def _emit_section(
    sec_number: str,
    sec_name: str,
    lines: list[str],
    parent_node: Node,
    origin_url: str,
) -> None:
    node_id = f"{parent_node.node_id}/section={sec_number}"
    citation = f"C.R.S. § {sec_number}"
    link = f"{origin_url}#{sec_number}"

    status = _check_reserved(sec_name)

    if status:
        section_node = Node(
            id=node_id,
            link=link,
            top_level_title=parent_node.top_level_title,
            node_type="content",
            level_classifier="section",
            number=sec_number,
            node_name=f"{sec_number}. {sec_name}",
            parent=parent_node.node_id,
            citation=citation,
            status=status,
        )
        insert_node(section_node, TABLE_NAME, ignore_duplicate=True, debug_mode=True)
        return

    node_text, addendum = _build_node_text(lines)

    section_node = Node(
        id=node_id,
        link=link,
        top_level_title=parent_node.top_level_title,
        node_type="content",
        level_classifier="section",
        number=sec_number,
        node_name=f"{sec_number}. {sec_name}",
        parent=parent_node.node_id,
        citation=citation,
        node_text=node_text,
        addendum=addendum,
    )
    insert_node(section_node, TABLE_NAME, ignore_duplicate=True, debug_mode=True)


# ---------------------------------------------------------------------------
# DOCX download + paragraph extraction
# ---------------------------------------------------------------------------

def _fetch_docx(title_slug: str) -> Optional[bytes]:
    """Download a title DOCX from the Wayback Machine with retries."""
    import time
    import requests

    url = f"{_WAYBACK_BASE}crs{_CRS_YEAR}-title-{title_slug}.docx"

    for attempt in range(1, _MAX_RETRIES + 1):
        try:
            from vaquill_pipeline.http_client import fetch_bytes
            body, _ct = fetch_bytes(url, timeout=60)
            return body
        except Exception:
            pass
        # Fallback: direct requests (no proxy needed for Wayback)
        try:
            r = requests.get(url, headers=_REQ_HEADERS, timeout=60, allow_redirects=True)
            if r.status_code == 200 and len(r.content) > 1000:
                return r.content
            print(
                f"[CO] title={title_slug} attempt={attempt} status={r.status_code}",
                flush=True,
            )
        except Exception as exc:
            print(
                f"[CO] title={title_slug} attempt={attempt} exc={exc!s:.80}",
                flush=True,
            )
        if attempt < _MAX_RETRIES:
            time.sleep(2 ** attempt)

    return None


def _extract_paragraphs(docx_bytes: bytes) -> list[str]:
    """Return list of non-empty, tab-stripped paragraph strings from DOCX."""
    from docx import Document  # type: ignore

    doc = Document(BytesIO(docx_bytes))
    result: list[str] = []
    for para in doc.paragraphs:
        text = para.text
        # Each paragraph starts with a leading tab in the CRS DOCX.
        text = text.lstrip("\t")
        text = _clean(text)
        if text:
            result.append(text)
    return result


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_title_header(paragraphs: list[str], top_level_title: str) -> tuple[str, int]:
    """
    Locate the "TITLE N" line in the paragraph list and derive the title name
    from the following line.

    Returns (title_name, index_of_TITLE_line).  Returns ("Title N", -1) if not found.
    """
    title_re = re.compile(rf"^TITLE\s+{re.escape(top_level_title.lstrip('0') or '0')}$")
    for i, text in enumerate(paragraphs):
        if title_re.match(text.strip()):
            # The next non-empty, non-note line is the name.
            name_parts: list[str] = []
            j = i + 1
            while j < len(paragraphs) and len(name_parts) < 3:
                nxt = paragraphs[j].strip()
                j += 1
                if not nxt:
                    continue
                # Stop collecting if we hit a structural keyword.
                if _is_structural(nxt):
                    break
                # Skip notes / cross references at title level.
                if nxt.startswith("Editor") or nxt.startswith("Cross ref"):
                    continue
                name_parts.append(nxt)
                if not nxt.isupper():
                    break
            title_name = f"Title {top_level_title.lstrip('0') or '0'}"
            if name_parts:
                title_name = f"Title {top_level_title.lstrip('0') or '0'} {' '.join(name_parts)}"
            return title_name, i
    return f"Title {top_level_title}", -1


def _is_structural(text: str) -> bool:
    """Return True if the text line is a top-level structural keyword."""
    return bool(
        re.match(r"^TITLE\s+\d", text)
        or re.match(r"^ARTICLE\s+\d", text)
        or re.match(r"^PART\s+\d", text)
        or _SEC_RE.match(text)
    )


def _extract_section_name(rest: str) -> str:
    """
    Extract the human-readable section name from the text following the
    section number and dot.

    E.g.:
        "Short title. Articles 1 to 13 of this title..."
        -> "Short title"
        "(Repealed)"
        -> "(Repealed)"
    """
    if not rest:
        return ""
    # If it starts with '(' it is a note/repealed marker.
    if rest.startswith("("):
        m = re.match(r"^(\([^)]+\))", rest)
        return m.group(1) if m else rest[:60]
    # Otherwise grab up to the first period that ends a short heading.
    # Section names are typically 1-10 words.
    m = re.match(r"^([^.]{3,80})\.", rest)
    if m:
        return m.group(1).strip()
    return rest[:80].strip()


def _build_node_text(
    lines: list[str],
) -> tuple[Optional[NodeText], Optional[Addendum]]:
    """Convert a list of raw section lines into NodeText + optional history Addendum."""
    if not lines:
        return None, None

    node_text = NodeText()
    history_parts: list[str] = []
    in_history = False

    _HISTORY_RE = re.compile(
        r"^Source:|^History:|^L\.\s*\d{2,4}:|^Editor",
        re.IGNORECASE,
    )

    for raw in lines:
        text = _clean(raw)
        if not text:
            continue
        if _HISTORY_RE.match(text):
            in_history = True
        if in_history:
            history_parts.append(text)
        else:
            node_text.add_paragraph(text=text)

    addendum: Optional[Addendum] = None
    if history_parts:
        addendum = Addendum()
        addendum.history = AddendumType(
            type="history", text=" ".join(history_parts)
        )

    return node_text, addendum


def _check_reserved(text: str) -> Optional[str]:
    """Return 'reserved' if text indicates the section was repealed/reserved."""
    lower = text.lower()
    for kw in _RESERVED_KEYWORDS:
        if kw in lower:
            return "reserved"
    return None


def _clean(raw: str) -> str:
    """Normalize whitespace and remove non-breaking spaces."""
    text = raw.replace("\xa0", " ").replace("‑", "-").replace("–", "-")
    text = re.sub(r"\s+", " ", text).strip()
    return text


# ---------------------------------------------------------------------------
# CLI entry
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    main()
