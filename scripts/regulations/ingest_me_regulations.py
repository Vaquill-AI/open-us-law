#!/usr/bin/env python3
"""Ingest the Code of Maine Rules (CMR) — Maine's state administrative rules.

Official source: Maine Secretary of State, Bureau of Corporations, Elections
& Commissions (CEC), Rulemaking Section.

    https://www.maine.gov/sos/cec/rules/                      (landing)
    https://www.maine.gov/sos/rulemaking/agency-rules         (CMR umbrella list)
    https://www.maine.gov/sos/rulemaking/agency-rules/{dept}  (chapters)

The CMR hierarchy is:

    Umbrella department  (e.g. "26 239 Department of Attorney General - General")
      └─ Sub-agency       (each h3 on a department page, sharing the same
                           umbrella-major code; different umbrella-minor codes)
           └─ Chapter      (one .doc / .docx / .pdf file per chapter, named
                            {minor}c{chapter}.doc[x], e.g. 239c001.doc)

A Maine rule's canonical citation is:

    {major}-{minor} Ch. {chapter}        e.g. "26-239 Ch. 1"

Maine ships chapters as Microsoft Word .docx (preferred), legacy Word .doc
(OLE binary), and a small handful of .pdf attachments. We handle all three:
docx via python-docx, .doc via olefile + cp1252 string extraction (the
WordDocument stream encodes English rule text directly as Latin-1 / Windows-
1252), and .pdf via pdfplumber.

Rule pages always end with a labeled metadata block:

    STATUTORY AUTHORITY: <citations>
    EFFECTIVE DATE:      <date>
    NON-SUBSTANTIVE CORRECTIONS:   <dates>
    REPEALED AND REPLACED:         <dates>
    AMENDED:                       <dates>

These fields are captured as structured metadata, not stripped silently.

Geo-restricted host; uses Webshare US-rotate + Mozilla UA + polite pacing.
corpus_type='state_regulation'.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
import re
import sys
import time
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path

import requests
from bs4 import BeautifulSoup

sys.stdout.reconfigure(line_buffering=True)  # type: ignore[attr-defined]

_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
DATA_DIR = Path(os.environ.get("OUT_DIR", "./data"))
OUT = DATA_DIR / "state_me_regulations.jsonl"

ME_HOST = "https://www.maine.gov"
ME_TOC = f"{ME_HOST}/sos/rulemaking/agency-rules"

_MOZ_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36"
)


def _load_env() -> None:
    env_path = _PROJECT_ROOT / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text().splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))


_load_env()


def _us_proxies() -> dict | None:
    user = os.environ.get("WEBSHARE_USERNAME", "")
    pwd = os.environ.get("WEBSHARE_PASSWORD", "")
    if not user or not pwd:
        return None
    import urllib.parse

    proxy_user = f"{user}-US-rotate"
    host = os.environ.get("WEBSHARE_PROXY_HOST", "p.webshare.io")
    port = os.environ.get("WEBSHARE_PROXY_PORT", "80")
    url = f"http://{urllib.parse.quote(proxy_user)}:{urllib.parse.quote(pwd)}@{host}:{port}"
    return {"http": url, "https": url}


SESSION = requests.Session()
SESSION.headers.update({"User-Agent": _MOZ_UA})


def fetch(url: str, retries: int = 5) -> str | None:
    """Fetch text. Same 429/502/503/504 backoff used in ingest_oh_regulations.py."""
    proxies = _us_proxies()
    for attempt in range(retries):
        try:
            r = SESSION.get(url, timeout=60, proxies=proxies, allow_redirects=True)
            if r.status_code == 200:
                return r.text
            if r.status_code == 429:
                time.sleep(2**attempt)
                continue
            if r.status_code in (502, 503, 504):
                time.sleep(2)
                continue
            return None
        except Exception:
            time.sleep(2)
    return None


def fetch_bytes(url: str, retries: int = 5) -> bytes | None:
    """Fetch raw bytes (for .docx / .doc / .pdf downloads)."""
    proxies = _us_proxies()
    for attempt in range(retries):
        try:
            r = SESSION.get(url, timeout=120, proxies=proxies, allow_redirects=True)
            if r.status_code == 200:
                return r.content
            if r.status_code == 429:
                time.sleep(2**attempt)
                continue
            if r.status_code in (502, 503, 504):
                time.sleep(2)
                continue
            return None
        except Exception:
            time.sleep(2)
    return None


# ---------------------------------------------------------------------------
# Discovery: agency-rules TOC -> department pages -> chapter file links
# ---------------------------------------------------------------------------

# Sub-agency h3: "<2-digit major> <2-3-digit minor> <name>"
_SUB_HEADING_RE = re.compile(r"^\s*(\d{2})\s+(\d{2,3})\b\s*(.+?)\s*$")
# Chapter file URL: .../{minor}c{chapter}*.doc or .docx (case-insensitive),
# also accept .pdf attachments. Maine sometimes pads chapter with zeros and
# appends suffixes like "%20-%20JUL%202025.docx" or "%20%28NEW%29.docx".
_CHAPTER_HREF_RE = re.compile(
    r"/(\d{2,3})c(\d{1,4})(?:[A-Za-z]?)(?:[^/]*?)\.(docx?|pdf)$",
    re.IGNORECASE,
)


def list_department_pages() -> list[str]:
    """Return the absolute URLs of every umbrella department page."""
    html = fetch(ME_TOC)
    if not html:
        raise RuntimeError("could not fetch ME agency-rules TOC")
    soup = BeautifulSoup(html, "html.parser")
    urls: list[str] = []
    seen: set[str] = set()
    for a in soup.find_all("a", href=True):
        h = a["href"]
        # Department pages live under /sos/rulemaking/agency-rules/<slug>.
        # Exclude the index itself and unrelated nav links.
        if "/rulemaking/agency-rules/" not in h:
            continue
        if h.rstrip("/").endswith("/agency-rules"):
            continue
        if "list-of-liaisons" in h:
            continue
        full = h if h.startswith("http") else f"{ME_HOST}{h}"
        if full in seen:
            continue
        seen.add(full)
        urls.append(full)
    return urls


@dataclass
class ChapterFile:
    umbrella_major: str  # e.g. "26"
    umbrella_minor: str  # e.g. "239"
    sub_agency_name: str  # e.g. "Department of Attorney General - General"
    chapter_num: str  # padded as it appears: "001", "104", "500"
    chapter_label: str  # e.g. "Ch. 1"
    chapter_title: str  # text following the link in the <li>
    file_url: str  # absolute URL to .doc / .docx / .pdf
    file_kind: str  # "docx" | "doc" | "pdf"
    department_url: str  # umbrella department page
    department_name: str  # human label for the umbrella department


def _walk_chapters_on_dept_page(html: str, dept_url: str) -> list[ChapterFile]:
    """Parse one umbrella department page into ChapterFile records.

    A page is a flat sequence of h3 sub-agency headings followed by <ul><li>
    rows with chapter file links. We walk the document order and assign each
    chapter file to the most recent h3.
    """
    soup = BeautifulSoup(html, "html.parser")
    main = (
        soup.find(id="block-sos-content") or soup.find("main") or soup.find(id="page-body") or soup
    )
    department_name = ""
    h1 = soup.find("h1")
    if h1:
        department_name = h1.get_text(" ", strip=True)

    out: list[ChapterFile] = []
    current_major = ""
    current_minor = ""
    current_sub_name = ""
    for el in main.descendants:
        if not hasattr(el, "name") or el.name is None:
            continue
        if el.name in ("h2", "h3", "h4"):
            txt = el.get_text(" ", strip=True)
            m = _SUB_HEADING_RE.match(txt)
            if m:
                current_major = m.group(1)
                current_minor = m.group(2)
                current_sub_name = m.group(3).strip()
            continue
        if el.name != "a":
            continue
        href = el.get("href") or ""
        m = _CHAPTER_HREF_RE.search(href.split("?", 1)[0])
        if not m:
            continue
        minor_in_href = m.group(1)
        chapter_num_raw = m.group(2)
        ext = m.group(3).lower()
        # If no h3 has been seen yet (some single-agency departments skip the
        # h3), fall back to inferring umbrella from the file name. ME numbers
        # the major code per department; without an h3 we cannot recover it,
        # so we mark major as "" and let the citation use only minor.
        minor = current_minor or minor_in_href
        # The <li> sometimes shows the label "Ch. 1" inside the link text and
        # the chapter title in the trailing text node. Capture both.
        link_text = el.get_text(" ", strip=True)
        # Pull "Ch. N" from link text if present, else from chapter_num_raw.
        chap_label_match = re.search(r"Ch\.\s*([\dA-Za-z\-]+)", link_text)
        chapter_label = (
            f"Ch. {chap_label_match.group(1)}"
            if chap_label_match
            else f"Ch. {int(chapter_num_raw)}"
        )
        # Title = sibling text after the <a>, before the next <a>/<li>.
        title_parts: list[str] = []
        sib = el.next_sibling
        while sib is not None:
            if hasattr(sib, "name") and sib.name in ("a", "li", "ul", "br"):
                break
            if hasattr(sib, "name") and sib.name is not None:
                title_parts.append(sib.get_text(" ", strip=True))
            else:
                title_parts.append(str(sib).strip())
            sib = sib.next_sibling
        title = " ".join(t for t in title_parts if t).strip(" -–—")
        full_url = href if href.startswith("http") else f"{ME_HOST}{href}"
        out.append(
            ChapterFile(
                umbrella_major=current_major,
                umbrella_minor=minor,
                sub_agency_name=current_sub_name,
                chapter_num=chapter_num_raw,
                chapter_label=chapter_label,
                chapter_title=title,
                file_url=full_url,
                file_kind="docx" if ext == "docx" else ("pdf" if ext == "pdf" else "doc"),
                department_url=dept_url,
                department_name=department_name,
            )
        )
    return out


def list_chapter_files() -> list[ChapterFile]:
    """Walk every umbrella department page and collect all chapter files."""
    pages = list_department_pages()
    all_files: list[ChapterFile] = []
    for p in pages:
        html = fetch(p)
        if not html:
            print(f"  ! could not fetch {p}", flush=True)
            continue
        files = _walk_chapters_on_dept_page(html, p)
        all_files.extend(files)
        time.sleep(0.4)
    return all_files


# ---------------------------------------------------------------------------
# Text extraction: .docx (python-docx), .doc (olefile), .pdf (pdfplumber)
# ---------------------------------------------------------------------------


def _extract_docx_text(blob: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(blob))
    parts: list[str] = []
    for p in doc.paragraphs:
        # Preserve leading tabs/indentation since rule structure (A., 1., (i))
        # is encoded via tab indent. Collapse internal runs of spaces only.
        t = p.text
        if t is None:
            continue
        t = t.rstrip()
        parts.append(t)
    # Also walk tables (some chapters embed schedules in tables).
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_doc_text(blob: bytes) -> str:
    """Best-effort .doc (OLE binary) text extraction.

    Reads the WordDocument stream and decodes as cp1252. Word's older binary
    format embeds English rule text as Latin-1 in this stream; we strip null
    and control bytes, anchor at the first "<major>-<minor>" citation header
    that every Maine rule starts with, then cut the trailing OLE chrome by
    detecting where letter density collapses.
    """
    import olefile

    ole = olefile.OleFileIO(io.BytesIO(blob))
    if not ole.exists("WordDocument"):
        return ""
    stream = ole.openstream("WordDocument").read()
    text = stream.decode("cp1252", errors="ignore")
    cleaned_chars: list[str] = []
    for c in text:
        if c == "\x00" or (ord(c) < 32 and c not in "\n\r\t"):
            cleaned_chars.append(" ")
        else:
            cleaned_chars.append(c)
    cleaned = "".join(cleaned_chars)
    # Anchor at the citation header (e.g. "26-239 DEPARTMENT OF...").
    m = re.search(r"\d{2}-\d{2,3}\s+[A-Z][A-Z ,&'\-/()]+", cleaned)
    if m:
        cleaned = cleaned[m.start() :]
    # Cut trailing OLE chrome: the first 200-char window with < 30 letters.
    win = 200
    cut = len(cleaned)
    i = 0
    while i + win <= len(cleaned):
        window = cleaned[i : i + win]
        letters = sum(1 for c in window if c.isalpha())
        if letters < 30:
            cut = i
            break
        i += 100
    cleaned = cleaned[:cut].rstrip()
    cleaned = re.sub(r"[ \t]{2,}", "  ", cleaned)
    cleaned = re.sub(r" *\n *", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _extract_pdf_text(blob: bytes) -> str:
    # pymupdf (fitz) — memory-safe alternative to pdfplumber for large-scale
    # scrapes; same output shape.
    import fitz  # pymupdf

    parts: list[str] = []
    with fitz.open(stream=blob, filetype="pdf") as pdf:
        for page in pdf:
            t = page.get_text() or ""
            if t.strip():
                parts.append(t.strip())
    return "\n\n".join(parts).strip()


def extract_chapter_text(blob: bytes, kind: str) -> str:
    if kind == "docx":
        return _extract_docx_text(blob)
    if kind == "doc":
        return _extract_doc_text(blob)
    if kind == "pdf":
        return _extract_pdf_text(blob)
    return ""


# ---------------------------------------------------------------------------
# Metadata parsing (statutory authority, effective date, history, etc.)
# ---------------------------------------------------------------------------


# Pull each labeled block from a Maine rule's trailing metadata section.
# Maine writes every label in ALL CAPS followed by ':'. The label may be at
# start of line OR immediately follow a sentence-ending punctuation (the .doc
# OLE extractor sometimes collapses newlines so "facilities.STATUTORY
# AUTHORITY:" is one run). To stay safe we anchor labels on:
#   (start | newline | "." | "?" | "!" | tab | run of 2+ spaces)
# AND require uppercase-ish label text followed by a colon. `REPEALED` excludes
# the longer `REPEALED AND REPLACED` form.
_META_LABELS: dict[str, str] = {
    "statutory_authority": r"STATUTORY\s+AUTHORITY",
    "effective_date": r"EFFECTIVE\s+DATE",
    "amended": r"AMENDED",
    "non_substantive_corrections": r"NON[\- ]SUBSTANTIVE\s+CORRECTIONS",
    "repealed_and_replaced": r"REPEALED\s+AND\s+REPLACED",
    "repealed": r"REPEALED(?!\s+AND\s+REPLACED)",
    "review_date": r"FIVE[\- ]YEAR\s+REVIEW",
    "conversion_check": r"WORD\s+VERSION\s+CONVERSION[^\n:]*",
}

# Extra stop-only labels (we don't want their values but we DO want to stop
# extraction when we hit them).
_STOP_ONLY_LABELS = [
    r"FISCAL\s+IMPACT\s+NOTE[^\n:]*",
    r"FISCAL\s+IMPACT",
    r"BASIS\s+STATEMENT",
    r"CONCISE\s+SUMMARY",
]


def _label_anchor(pat: str) -> str:
    """Match label only at a non-letter boundary (start-of-string or any non-
    alphabetic char preceding it). The label MUST be immediately followed by
    ':'. Maine writes every metadata label as ALL-CAPS + ':', so requiring
    `[A-Za-z]` not preceding lets us match both the docx case (label at line
    start) and the .doc case (label runs onto the same line as the previous
    value).
    """
    return rf"(?:^|(?<=[^A-Za-z])){pat}\s*:\s*"


def _stop_pattern() -> str:
    parts = [_label_anchor(p) for p in _META_LABELS.values()]
    parts.extend(_label_anchor(p) for p in _STOP_ONLY_LABELS)
    return r"(?:" + "|".join(parts) + r")"


def _extract_field(text: str, label_pat: str) -> str:
    label_re = re.compile(_label_anchor(label_pat), re.IGNORECASE | re.MULTILINE)
    m = label_re.search(text)
    if not m:
        return ""
    start = m.end()
    stop_re = re.compile(_stop_pattern(), re.IGNORECASE | re.MULTILINE)
    n = stop_re.search(text, start)
    end = n.start() if n else len(text)
    val = text[start:end].strip(" \t:\n\r")
    val = re.sub(r"\s+", " ", val)
    val = re.sub(r"\s*page\s*PAGE\s*$", "", val, flags=re.IGNORECASE).strip()
    return val


def parse_metadata_block(text: str) -> dict[str, str]:
    out: dict[str, str] = {}
    for key, pat in _META_LABELS.items():
        v = _extract_field(text, pat)
        if v:
            out[key] = v
    return out


def _strip_metadata_tail(text: str) -> str:
    """Return the body with the trailing labeled metadata block removed."""
    m = re.search(_stop_pattern(), text, re.IGNORECASE | re.MULTILINE)
    if m and m.start() > 200:  # only cut if there's a real body before it
        return text[: m.start()].rstrip()
    return text


# ---------------------------------------------------------------------------
# Chunk emission
# ---------------------------------------------------------------------------


@dataclass
class Rule:
    cf: ChapterFile
    raw_text: str
    chapter_heading: str = ""  # e.g. "Chapter 3: STANDARDS FOR ..."
    statutory_authority: str = ""
    effective_date: str = ""
    last_amended_date: str = ""
    prior_effective_dates: str = ""
    repealed_and_replaced: str = ""
    review_date: str = ""
    history_notes: str = ""


def _sha1(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8")).hexdigest()


def _point_id(act_id: str, idx: int, text: str) -> str:
    seed = f"{act_id}::{idx}::{_sha1(text)[:12]}"
    return str(uuid.UUID(hashlib.md5(seed.encode()).hexdigest()))


def _safe(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_")


def _chapter_heading_from_body(text: str) -> str:
    """Pull the inline chapter heading. Maine documents always start with
    "<major>-<minor>\\tDEPARTMENT NAME\\nChapter N: TITLE".
    """
    m = re.search(r"Chapter\s+[\dA-Za-z\-]+\s*:?\s*[^\n]+", text)
    if m:
        head = m.group(0)
        return re.sub(r"\s+", " ", head).strip().rstrip(".")
    return ""


def _drop_header_lines(text: str) -> str:
    """Remove the umbrella header + chapter heading from the body (they're
    already captured in metadata), so the embedded text starts with the rule's
    substantive opening (usually "SUMMARY:" or "§ 1").

    Bounded so we never delete past the start of the substantive body. The
    .doc extractor collapses some newlines, so the headings can run into the
    next line; we stop at the first occurrence of any common rule-body anchor
    ("SUMMARY:", "Section", "§", "PURPOSE:", or an arabic-section "1.").
    """
    out = text
    # Find where the substantive body starts.
    anchor_re = re.compile(
        r"(?:\n\s*)?(?:SUMMARY|PURPOSE|SCOPE|AUTHORITY|DEFINITIONS|SECTION|PREAMBLE)\s*[:\.]"
        r"|\n\s*§\s*\d"
        r"|\n\s*Section\s+\d"
        r"|\n\s*[IVX]+\.\s+[A-Z]"
        r"|\n\s*1\.\s+[A-Z]",
        re.IGNORECASE,
    )
    m = anchor_re.search(out)
    if m:
        out = out[m.start() :].lstrip("\n")
        # Drop a leading newline-only run.
        return out.lstrip()
    # Fallback (no recognised anchor): remove just the first two header lines.
    out = re.sub(r"^\s*\d{2}-\d{2,3}\s+[A-Z][A-Z ,&'\-/()]+\n?", "", out, count=1)
    out = re.sub(r"^\s*Chapter\s+[\dA-Za-z\-]+\s*:?\s*[^\n]+\n?", "", out, count=1)
    return out.lstrip()


def build_rule(cf: ChapterFile, raw_text: str) -> Rule | None:
    if not raw_text or len(raw_text) < 60:
        return None
    chapter_heading = _chapter_heading_from_body(raw_text)
    meta = parse_metadata_block(raw_text)
    body = _strip_metadata_tail(raw_text)
    body = _drop_header_lines(body)
    body = re.sub(r"\s+\n", "\n", body).strip()
    if len(body) < 60:
        return None
    return Rule(
        cf=cf,
        raw_text=body,
        chapter_heading=chapter_heading,
        statutory_authority=meta.get("statutory_authority", ""),
        effective_date=meta.get("effective_date", ""),
        last_amended_date=meta.get("amended", ""),
        prior_effective_dates=meta.get("non_substantive_corrections", ""),
        repealed_and_replaced=meta.get("repealed_and_replaced", ""),
        review_date=meta.get("review_date", ""),
        history_notes="; ".join(
            v for v in (meta.get("repealed", ""), meta.get("conversion_check", "")) if v
        ),
    )


def to_chunk_record(rule: Rule) -> dict:
    cf = rule.cf
    # act_id is a sanitized metadata identifier.
    safe_major = cf.umbrella_major or "00"
    safe_minor = cf.umbrella_minor or "000"
    safe_chap = _safe(cf.chapter_num.lstrip("0") or cf.chapter_num)
    act_id = f"STATE_ME_CMR_{safe_major}_{safe_minor}_{safe_chap}"
    # ME citation: "26-239 Ch. 1"
    chapter_num_clean = cf.chapter_num.lstrip("0") or "0"
    if cf.umbrella_major:
        citation = f"{cf.umbrella_major}-{cf.umbrella_minor} Ch. {chapter_num_clean}"
    else:
        citation = f"{cf.umbrella_minor} Ch. {chapter_num_clean}"
    chapter_title_clean = cf.chapter_title.strip() or rule.chapter_heading
    display_title = (
        f"{cf.chapter_label}{(': ' + chapter_title_clean) if chapter_title_clean else ''}"
    )

    text = rule.raw_text
    # Embed header surfaces the regulation-specific metadata so it stays
    # retrievable, not just stored. Mirrors the OH pattern.
    meta_lines: list[str] = []
    if rule.effective_date:
        meta_lines.append(f"Effective: {rule.effective_date}")
    if rule.last_amended_date:
        meta_lines.append(f"Amended: {rule.last_amended_date}")
    if rule.statutory_authority:
        meta_lines.append(f"Statutory Authority: {rule.statutory_authority}")
    if rule.review_date:
        meta_lines.append(f"Review: {rule.review_date}")
    meta_header = ("\n" + "\n".join(meta_lines)) if meta_lines else ""
    sub_agency_label = cf.sub_agency_name or cf.department_name
    text_for_embedding = (
        f"Regulation: Code of Maine Rules | US | Maine | In Force\n"
        f"{cf.umbrella_major}-{cf.umbrella_minor} {sub_agency_label}\n"
        f"{display_title}{meta_header}\n\n{text}"
    )

    md = {
        "act_id": act_id,
        "corpus_type": "state_regulation",
        "category": "state_regulation",
        "document_type": "regulation",
        "jurisdiction": "US",
        "country_code": "US",
        "state": "me",
        "title_number": None,
        "title_name": f"Code of Maine Rules — {sub_agency_label}",
        "title": "Code of Maine Rules",
        "title_code": "regs_me",
        "top_level_title": "regs-me",
        "chapter": cf.chapter_num,
        "chapter_name": chapter_title_clean,
        "section_number": citation,
        "section_title": display_title,
        "year": 2026,
        "act_status": "in_force",
        "renumbered_to": "",
        "transferred_to": "",
        "level_classifier": "section",
        # --- regulation-specific rich metadata (captured, not discarded) ---
        "effective_date": rule.effective_date or None,
        "statutory_authority": rule.statutory_authority or None,
        "rule_amplifies": None,  # ME does not publish a separate "amplifies" field
        "promulgated_under": None,
        "prior_effective_dates": rule.prior_effective_dates or None,
        "review_date": rule.review_date or None,
        "last_amended_date": rule.last_amended_date or None,
        "repealed_and_replaced": rule.repealed_and_replaced or None,
        "history_notes": rule.history_notes or None,
        "issuing_agency": sub_agency_label or None,
        "issuing_agency_code": f"{cf.umbrella_major}-{cf.umbrella_minor}".strip("-"),
        "umbrella_major": cf.umbrella_major or None,
        "umbrella_minor": cf.umbrella_minor or None,
        "department_name": cf.department_name or None,
        "source_format": cf.file_kind,
        "citation": citation,
        "citation_short": citation,
        "display_label": citation,
        "display_title": display_title,
        "display_path": (
            f"Code of Maine Rules / {cf.umbrella_major}-{cf.umbrella_minor} "
            f"{sub_agency_label} / {cf.chapter_label}"
        ),
        "breadcrumb": [
            {
                "type": "umbrella",
                "num": cf.umbrella_major,
                "label": f"Umbrella {cf.umbrella_major}",
                "name": cf.department_name,
            },
            {
                "type": "agency",
                "num": cf.umbrella_minor,
                "label": f"{cf.umbrella_major}-{cf.umbrella_minor}",
                "name": sub_agency_label,
            },
            {
                "type": "chapter",
                "num": cf.chapter_num,
                "label": cf.chapter_label,
                "name": chapter_title_clean,
            },
        ],
        "sort_key": act_id,
        "word_count": len(text.split()),
        "subsection_count": 0,
        "subsection_letters": [],
        "numbered_paragraph_count": 0,
        "cross_references_count": 0,
        "cross_references_usc": [],
        "cross_references_cfr": [],
        "amendment_years": [],
        "amendments_count": 0,
        "last_amended_year": None,
        "public_laws_referenced": [],
        "public_laws_count": 0,
        "source_url": cf.file_url,
        "parent_id": (f"us/me/regulations/agency={cf.umbrella_major}-{cf.umbrella_minor}"),
        "raw_node_id": (
            f"us/me/regulations/agency={cf.umbrella_major}-{cf.umbrella_minor}"
            f"/chapter={cf.chapter_num}"
        ),
        "parent_chunk_id": _point_id(act_id, -1, text),
        "full_text_sha1": _sha1(text),
    }
    return {
        "point_id": _point_id(act_id, 0, text),
        "text_for_embedding": text_for_embedding,
        "raw_text": text,
        "metadata": md,
    }


# ---------------------------------------------------------------------------
# Crawl
# ---------------------------------------------------------------------------


def process_chapter_file(cf: ChapterFile) -> list[Rule]:
    blob = fetch_bytes(cf.file_url)
    if not blob:
        return []
    try:
        text = extract_chapter_text(blob, cf.file_kind)
    except Exception as e:
        print(f"  ! extract failed {cf.file_url}: {e}", flush=True)
        return []
    rule = build_rule(cf, text)
    if rule is None:
        return []
    return [rule]


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--umbrellas",
        default="",
        help=(
            "Comma-separated umbrella-minor codes to restrict the crawl "
            "(e.g. '239,550'). Default: all."
        ),
    )
    ap.add_argument(
        "--department-slugs",
        default="",
        help=(
            "Comma-separated department slugs (e.g. "
            "'department-attorney-general-rules'). Default: all."
        ),
    )
    ap.add_argument(
        "--max-files",
        type=int,
        default=0,
        help="Cap total files processed (0 = no cap; for smoke tests).",
    )
    ap.add_argument("--workers", type=int, default=3)
    ap.add_argument("--delay", type=float, default=0.6)
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args()

    DATA_DIR.mkdir(parents=True, exist_ok=True)

    print(f"[CMR] discovering departments from {ME_TOC}", flush=True)
    pages = list_department_pages()
    if args.department_slugs:
        wanted = {s.strip() for s in args.department_slugs.split(",") if s.strip()}
        pages = [p for p in pages if any(p.rstrip("/").endswith(w) for w in wanted)]
    print(f"[CMR] {len(pages)} umbrella department pages", flush=True)

    # Phase 1: gather all chapter files from each department.
    all_files: list[ChapterFile] = []
    for p in pages:
        html = fetch(p)
        if not html:
            print(f"  ! could not fetch {p}", flush=True)
            continue
        files = _walk_chapters_on_dept_page(html, p)
        all_files.extend(files)
        print(f"  [{p.rsplit('/', 1)[-1]}] {len(files)} chapter files", flush=True)
        time.sleep(args.delay)
    if args.umbrellas:
        wanted = {u.strip() for u in args.umbrellas.split(",") if u.strip()}
        all_files = [f for f in all_files if f.umbrella_minor in wanted]
    if args.max_files > 0:
        all_files = all_files[: args.max_files]
    print(f"\n[CMR] {len(all_files)} chapter files to fetch", flush=True)
    if args.dry_run:
        return 0

    # Phase 2: download + parse + emit chunks. Stream-write per rule so the
    # process doesn't buffer all chunks in memory (earlier accumulate pattern
    # was unbounded).
    seen: set[str] = set()
    if OUT.exists():
        with open(OUT) as fh:
            for line in fh:
                try:
                    seen.add(json.loads(line)["point_id"])
                except Exception:
                    pass
    parsed = 0
    written = 0
    t0 = time.time()
    with open(OUT, "a") as out_fh, ThreadPoolExecutor(max_workers=args.workers) as ex:
        futures = {ex.submit(process_chapter_file, cf): cf for cf in all_files}
        done = 0
        for fut in as_completed(futures):
            done += 1  # noqa: SIM113 (counter over as_completed, not a positional index)
            cf = futures[fut]
            try:
                rules = fut.result()
                for rule in rules:
                    rec = to_chunk_record(rule)
                    parsed += 1
                    if rec["point_id"] not in seen:
                        out_fh.write(json.dumps(rec, ensure_ascii=False) + "\n")
                        out_fh.flush()
                        seen.add(rec["point_id"])
                        written += 1
            except Exception as e:
                print(f"  ! chapter failed {cf.file_url}: {e}", flush=True)
            if done % 50 == 0 or done == len(all_files):
                rate = parsed / max(time.time() - t0, 0.1)
                print(
                    f"  ... {done:>5}/{len(all_files)} files, "
                    f"{parsed:>6} rules, {written} written, {rate:.1f}/s",
                    flush=True,
                )
            time.sleep(args.delay / max(args.workers, 1))

    print(
        f"\n=== Done: parsed={parsed:,}, new={written:,}, elapsed={time.time() - t0:.1f}s ===",
        flush=True,
    )
    print(f"JSONL: {OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
